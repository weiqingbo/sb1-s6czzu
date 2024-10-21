import pythoncom
from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
import requests
import json
from docx import Document
from docx.shared import Pt, Inches
import os
import shutil
from docx2pdf import convert
from docx.oxml.ns import qn


#
# from flask_sqlalchemy import SQLAlchemy
# from sqlalchemy.engine.url import URL
#
# # 创建Flask应用实例
# app = Flask(__name__)
#
# # 配置数据库URI
# DATABASE_URI = URL(
#     drivername='mysql+pymysql',  # 使用pymysql作为MySQL的驱动
#     username='your_mysql_username',  # 你的MySQL用户名
#     password='your_mysql_password',  # 你的MySQL密码
#     host='your_mysql_host',  # MySQL服务器地址，例如'localhost'或IP地址
#     port=3306,  # MySQL服务器端口，默认是3306
#     database='your_database_name'  # 你要连接的数据库名
# )
#
# # 将数据库URI配置到Flask-SQLAlchemy中
# app.config['SQLALCHEMY_DATABASE_URI'] = DATABASE_URI
# app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False  # 关闭信号跟踪，提高性能
#
# # 初始化SQLAlchemy实例
# db = SQLAlchemy(app)




app = Flask(__name__)
CORS(app)


# 创建需要的目录
if not os.path.exists('img'):
    os.makedirs('img')
if not os.path.exists('document'):
    os.makedirs('document')


def clear_directory(directory):
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path): #检查路径是否是文件或符号链接。
                os.unlink(file_path)
            elif os.path.isdir(file_path):  #检查路径是否是目录。
                shutil.rmtree(file_path)
        except Exception as e:
            print(f'Failed to delete {file_path}. Reason: {e}')


def set_font(run, name, size):
    run.font.name = name
    run.font.size = Pt(size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), name)


def get_access_token():
    url = "https://aip.baidubce.com/oauth/2.0/token?grant_type=client_credentials&client_id=s4ZmMkYlKsS1u9hh95kk8Crm&client_secret=ykp9MrTbh4XTfhgBAK15Zh3jG8w7IvN6"
    payload = json.dumps("")
    headers = {
        'Content-Type': 'application/json',
        'Accept': 'application/json'
    }

    response = requests.request("POST", url, headers=headers, data=payload)
    return response.json().get("access_token")


def call_ernie_bot(question):
    url = "https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/completions_pro?access_token=" + get_access_token()

    payload = json.dumps({
        "messages": [
            {
                "role": "user",
                "content": f"{question}"
            }
        ]
    })
    headers = {
        'Content-Type': 'application/json'
    }

    response = requests.request("POST", url, headers=headers, data=payload)
    return response.json()


resume_content = ''

def resume_fun(content):
    global resume_content
    resume_content = content

image_path = ''

def image_fun(content):
    global image_path
    image_path = content


@app.route('/generate_resume', methods=['POST'])
def generate_resume():
    # 清空 document 和 img 目录
    clear_directory('document/before')
    clear_directory('img')

    data = request.form
    name = data.get("name")
    email = data.get("email")
    phone = data.get("phone")
    location = data.get("location")
    hobby=data.get("hobby")
    job_position = data.get("job_position")
    specialty = data.get("specialty")


    file = request.files['avatar']

    if file:
        image_path = os.path.join('img', file.filename)
        image_fun(image_path)
        file.save(image_path)
    else:
        image_path = "img\头像.jpg"

    print("---简历制作中---")
    # 使用ERNIE Bot API生成简历内容
    profile_prompt = f"为{name}生成一个与{job_position}岗位相关的简短自我介绍，然后这个是兴趣爱好：{hobby}。注意你的回答中不要出现名字，自我介绍记得用第一人称。注意你的最后回答中不需要‘出现感谢大家的聆听，期待与您进一步交流。’或者‘谢谢大家’之类的的字眼；你的开头回答也不需要出现“大家好，”的介绍；直接给出自我介绍就行"
    profile_response = call_ernie_bot(profile_prompt)
    profile =profile_response.get("result", "未能生成自我介绍")
    profile = profile.split('\n')

    # 移除空字符串
    profile = [item for item in profile if item]


    experience_prompt = f"请为{name}生成3个工作经历，重点突出与{job_position}岗位相关的工作经历。注意你的回答中不要出现名字，直接给出{job_position}相关的工作经历就行，不要给我回答无关的字符，例如：*、+等等。请你根据以下的模板格式来生成工作经历，模板为：" \
                        f" 一、XXXX 科技公司 Python 网络爬虫工程师 2022 年 9 月至今" \
                        f"（1）" \
                        f"（2）" \
                        f"（3）" \
                        f"（4）" \
                        f"请注意每一个工作经历都要严格按照上面的模板给我生成，例如：一、 XXXX 科技公司 Python 网络爬虫工程师（xxxx 年 x 月至今）  再（1）（2）（3）（4） 二、XXXX 数据公司 Python 网络爬虫工程师（XXXX 年 X 月 至 XXXX 年 X 月）  再（1）（2）（3）（4）  以此类推"

    experience_response = call_ernie_bot(experience_prompt)
    experience = experience_response.get("result", "未能生成工作经历")
    experience = experience.split('\n')

    experience_cleaned_list = []
    for item in experience:
        cleaned_item = item.replace('*', '').replace('-', '').replace('+', '').strip()
        if cleaned_item:
            experience_cleaned_list.append(cleaned_item)

    education_prompt = f"请为{name}生成4个项目经历，重点突出与{job_position}岗位相关的项目经历。注意你的回答中不要出现名字，直接给出{job_position}相关的项目经历就行，不要给我回答无关的字符，例如：*、+等等。请你根据以下的模板格式来生成项目经历，模板为：" \
                       f" 一、 xxxxxxxx项目" \
                       f"（1）" \
                       f"（2）" \
                       f"（3）" \
                       f"（4）" \
                       f"请注意每一个项目经历都要严格按照上面的模板给我生成，例如：一、 xxxxxxxx项目  再（1）（2）（3）（4） 二、xxxxxxxx  再（1）（2）（3）（4）  以此类推"

    education_response = call_ernie_bot(education_prompt)
    education = education_response.get("result", "未能生成项目经历")
    education = education.split('\n')

    education_cleaned_list = []
    for item in education:
        cleaned_item = item.replace('*', '').replace('-', '').replace('+', '').strip()
        if cleaned_item:
            education_cleaned_list.append(cleaned_item)

    skills_prompt = f"{name}的专业为{specialty}，请列举4个有关{specialty}的专业技能，特别是与{job_position}相关的专业技能。注意你的回答中不要出现名字，直接给出专业技能就行，不用再生成其他的内容。" \
                    f"不要再生成类似于以下的语句，例如：（这里列举的专业技能均与xxx相关，并且在一定程度上与 xxx有关联。" \
                    f"注意，除上述专业技能以外，还可能有其他符合要求的答案。在具体应用中，这些专业技" \
                    f"能的掌握程度和运用方式会因实际情况而有所不同。如果需要更详细的解释或指导，请咨询专家。"

    skills_response = call_ernie_bot(skills_prompt)
    skills = skills_response.get("result", "未能生成技能列表")
    skills = skills.split('\n')

    # 移除空字符串
    skills = [item for item in skills if item]


    # 创建Word文档并添加内容
    document = Document()

    if image_path and os.path.exists(image_path):
        document.add_picture(image_path, width=Inches(1.0))

    document.add_heading(name.upper(), level=0)  #添加标题


    email_paragraph=document.add_paragraph()
    email_run=email_paragraph.add_run(email)
    set_font(email_run, 'Arial', 12)
    email_run.bold = False  # 确保不加粗

    phone_paragraph = document.add_paragraph()  #添加段落
    phone_run = phone_paragraph.add_run(phone)  #在段落中添加文本
    set_font(phone_run, 'Arial', 12)
    phone_run.bold = False  # 确保不加粗

    location_paragraph = document.add_paragraph()
    location_run = location_paragraph.add_run(location)
    set_font(location_run, 'Arial', 12)
    location_run.bold = False  # 确保不加粗


    document.add_heading("自我介绍", level=1)
    for i in profile:
        i="       "+i
        profile_paragraph = document.add_paragraph()
        profile_run = profile_paragraph.add_run(i)
        set_font(profile_run, 'Arial', 12)
        profile_run.bold = False  # 确保不加粗

    document.add_heading("工作经历", level=1)
    ex_index = 0
    for i in experience_cleaned_list:
        para = document.add_paragraph()
        run = para.add_run(i)
        set_font(run, 'Arial', 12)
        if ex_index == 0 or ex_index == 5 or ex_index == 10:
            run.bold = True
        ex_index += 1

    document.add_heading("项目经历", level=1)
    index = 0
    for i in education_cleaned_list:
        para = document.add_paragraph()
        run = para.add_run(i)
        set_font(run, 'Arial', 12)
        if index == 0 or index == 5 or index == 10 or index == 15:
            run.bold = True
        index += 1

    document.add_heading("专业技能", level=1)
    for i in skills:
        para = document.add_paragraph()
        run = para.add_run(i)
        set_font(run, 'Arial', 12)
        run.bold = False  # 确保不加粗

    docx_file = os.path.join('document/before', "简历.docx")
    document.save(docx_file)

    pythoncom.CoInitialize()  # 初始化 COM 库
    try:
        pdf_file = os.path.join('document/before', "简历.pdf")
        convert(docx_file, pdf_file)
    finally:
        pythoncom.CoUninitialize()  # 取消初始化 COM 库

    resume_content = {
        "profile": profile,
        "experience": experience_cleaned_list,
        "education": education_cleaned_list,
        "skills": skills
    }

    resume_fun(resume_content)
    return send_file(pdf_file, as_attachment=True)  #使用 send_file() 发送文件，并设置为附件

@app.route('/edit_resume', methods=['POST'])
def edit_resume():
    try:
        clear_directory('document/after')
        data = request.json
        name = data.get("name")
        email = data.get("email")
        phone = data.get("phone")
        location = data.get("location")
        edit_content = data.get("content", "")

        # 获取原来生成的简历内容
        original_profile = resume_content.get("profile")

        original_experience = resume_content.get("experience")
        gpt_original_experience = [line for line in original_experience if line]
        gpt_original_experience = '\n'.join(gpt_original_experience)

        original_education = resume_content.get("education")
        gpt_original_education = [line for line in original_education if line]
        gpt_original_education = '\n'.join(gpt_original_education)

        original_skills = resume_content.get("skills")

        profile, experience_cleaned_list, education_cleaned_list, skills_cleaned_list = original_profile, original_experience, original_education, original_skills

        if "自我介绍" in edit_content:
            profile_prompt = f"我有一段简短的自我介绍需要完善和修改。这是简短的自我介绍：{original_profile};这是修改自我介绍的需求：{edit_content.split('自我介绍部分，')[1].split('。')[0]},请你根据需求，来完善和修改自我介绍。注意你的最后回答中不需要‘出现感谢大家的聆听，期待与您进一步交流。’或者‘谢谢大家’之类的的字眼；你的开头回答也不需要出现“大家好，”的介绍；直接给出修改后的自我介绍就行"

            profile_response = call_ernie_bot(profile_prompt)
            profile = profile_response.get("result", original_profile).split('\n')

            profile = [s for s in profile if s]

        if "工作经历" in edit_content:
            experience_prompt = f"我有一段工作经历的文本需要完善和修改。这是我的工作经历:{gpt_original_experience};这是修改工作经历的需求:{edit_content.split('工作经历部分，')[1].split('。')[0]},请你根据需求，来完善和修改工作经历。请你严格按照原本工作经历的模板来生成修改后的工作经历，直接给出修改后的工作经历就行，不用再描述一下其他东西"
            experience_response = call_ernie_bot(experience_prompt)
            experience = experience_response.get("result", original_experience).split('\n')

            experience_cleaned_list = []
            for item in experience:
                cleaned_item = item.replace('*', '').replace('-', '').replace('+', '').strip()
                if cleaned_item:
                    experience_cleaned_list.append(cleaned_item)

        if "项目经历" in edit_content:
            education_prompt = f"我有一段项目经历的文本需要完善和修改。这是我的项目经历:{gpt_original_education};这是修改项目经历的需求:{edit_content.split('项目经历部分，')[1].split('。')[0]},请你根据需求，来完善和修改项目经历。请你严格按照原本项目经历的模板来生成修改后的工作经历，直接给出修改后的项目经历就行，不用再描述一下其他东西"
            education_response = call_ernie_bot(education_prompt)
            education = education_response.get("result", original_education).split('\n')

            education_cleaned_list = []
            for item in education:
                cleaned_item = item.replace('*', '').replace('-', '').replace('+', '').strip()
                if cleaned_item:
                    education_cleaned_list.append(cleaned_item)

        if "专业技能" in edit_content:
            skills_prompt = f"我有一段专业技能的文本需要完善和修改。这是我的专业技能:{original_skills};这是修改专业技能的需求:{edit_content.split('专业技能部分，')[1].split('。')[0]},请你根据需求，来完善和修改专业技能。请你严格按照原本专业技能的模板来生成修改后的专业技能，直接给出修改后的专业技能就行，不用再描述一下其他东西"
            skills_response = call_ernie_bot(skills_prompt)
            skills = skills_response.get("result", original_skills)
            skills = skills.split('\n')

            skills_cleaned_list = []
            for item in skills:
                cleaned_item = item.replace('*', '').replace('-', '').replace('+', '').strip()
                if cleaned_item:
                    skills_cleaned_list.append(cleaned_item)

        # 创建Word文档并添加内容
        document = Document()

        if image_path and os.path.exists(image_path):
            document.add_picture(image_path, width=Inches(1.0))

        def set_font(run, name, size):
            run.font.name = name
            run.font.size = Pt(size)
            r = run._element
            rPr = r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), name)

        document.add_heading(name.upper(), level=0)

        email_paragraph = document.add_paragraph()
        email_run = email_paragraph.add_run(email)
        set_font(email_run, 'Arial', 12)
        email_run.bold = False  # 确保不加粗

        phone_paragraph = document.add_paragraph()
        phone_run = phone_paragraph.add_run(phone)
        set_font(phone_run, 'Arial', 12)
        phone_run.bold = False  # 确保不加粗

        location_paragraph = document.add_paragraph()
        location_run = location_paragraph.add_run(location)
        set_font(location_run, 'Arial', 12)
        location_run.bold = False  # 确保不加粗

        document.add_heading("自我介绍", level=1)
        for i in profile:
            i = "       " + i
            profile_paragraph = document.add_paragraph()
            profile_run = profile_paragraph.add_run(i)
            set_font(profile_run, 'Arial', 12)
            profile_run.bold = False  # 确保不加粗

        document.add_heading("工作经历", level=1)
        ex_index = 0
        for i in experience_cleaned_list:
            para = document.add_paragraph()
            run = para.add_run(i)
            set_font(run, 'Arial', 12)
            if ex_index == 0 or ex_index == 5 or ex_index == 10:
                run.bold = True
            ex_index += 1

        document.add_heading("项目经历", level=1)
        index = 0
        for i in education_cleaned_list:
            para = document.add_paragraph()
            run = para.add_run(i)
            set_font(run, 'Arial', 12)
            if index == 0 or index == 5 or index == 10 or index == 15:
                run.bold = True
            index += 1

        document.add_heading("专业技能", level=1)
        index = 0
        if len(skills_cleaned_list)==4:
            for i in skills_cleaned_list:
                para = document.add_paragraph()
                run = para.add_run(i)
                set_font(run, 'Arial', 12)
                run.bold = False  # 确保不加粗
        else:
            for i in skills_cleaned_list:
                para = document.add_paragraph()
                run = para.add_run(i)
                set_font(run, 'Arial', 12)
                if index == 0 or index == 4 or index == 8 or index == 12:
                    run.bold = True
                index += 1

        docx_file = os.path.join('document/after', f"修改简历.docx")
        document.save(docx_file)

        pythoncom.CoInitialize()  # 初始化 COM 库
        try:
            pdf_file = os.path.join('document/after', f"修改简历.pdf")
            convert(docx_file, pdf_file)
        finally:
            pythoncom.CoUninitialize()  # 取消初始化 COM 库

        return send_file(pdf_file, as_attachment=True)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == '__main__':
    app.run(host="0.0.0.0", port=5000, debug=True)


